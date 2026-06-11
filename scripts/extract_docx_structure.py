"""Extract a DOCX manuscript structure into Markdown.

Usage:
    python scripts/extract_docx_structure.py manuscript.docx --out structure.md
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def is_caption(text: str) -> bool:
    stripped = text.strip()
    return stripped.lower().startswith(("figure ", "fig. ", "table "))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out", type=Path, default=None)
    args = parser.parse_args()

    doc = Document(args.docx)
    lines: list[str] = []
    lines.append(f"# DOCX Structure: {args.docx.name}")
    lines.append("")
    lines.append(f"- Paragraphs: {len(doc.paragraphs)}")
    lines.append(f"- Tables: {len(doc.tables)}")
    lines.append("")

    lines.append("## Headings and Captions")
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else ""
        if not text:
            continue
        if style.lower().startswith("heading") or is_caption(text) or text.lower() in {
            "abstract",
            "keywords",
            "references",
            "introduction",
            "discussion",
            "conclusion",
            "conclusions",
        }:
            lines.append(f"- P{i}: `{style}` - {text}")

    lines.append("")
    lines.append("## Tables")
    for idx, table in enumerate(doc.tables, 1):
        rows = len(table.rows)
        cols = len(table.columns)
        first = " | ".join(cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells) if rows else ""
        lines.append(f"- Table object {idx}: {rows} rows x {cols} columns")
        if first:
            lines.append(f"  - Header: {first}")

    output = "\n".join(lines)
    if args.out:
        args.out.write_text(output, encoding="utf-8")
    else:
        print(output)


if __name__ == "__main__":
    main()

