"""Apply simple three-line-table formatting to all DOCX tables.

This script removes vertical borders and most internal horizontal borders,
leaving the top rule, header-bottom rule, and bottom rule.

Usage:
    python scripts/format_three_line_tables.py manuscript.docx --out manuscript_tables.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        data = kwargs.get(edge)
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        if data is None:
            element.set(qn("w:val"), "nil")
        else:
            for key, value in data.items():
                element.set(qn(f"w:{key}"), str(value))


def format_table(table) -> None:
    nrows = len(table.rows)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(9)
        if r_idx == 0:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top={"val": "single", "sz": "12", "space": "0", "color": "000000"},
                    bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"},
                    left=None,
                    right=None,
                    insideH=None,
                    insideV=None,
                )
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        if r_idx == nrows - 1:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    bottom={"val": "single", "sz": "12", "space": "0", "color": "000000"},
                    left=None,
                    right=None,
                    insideH=None,
                    insideV=None,
                )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()

    doc = Document(args.docx)
    for table in doc.tables:
        format_table(table)
    doc.save(args.out)
    print(args.out)


if __name__ == "__main__":
    main()

